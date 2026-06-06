from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "docs/analisis-arquitectura-live-control.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
      set_cell_text(header_cells[index], header, bold=True)
      set_cell_shading(header_cells[index], "E8EEF5")
      header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      if widths:
          header_cells[index].width = Inches(widths[index])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cells[index].width = Inches(widths[index])
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = RGBColor(46, 116, 181)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(46, 116, 181)
            run.font.size = Pt(13)
        else:
            run.font.color.rgb = RGBColor(31, 77, 120)
            run.font.size = Pt(12)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Analisis tecnico del proyecto Live Control App")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run(
        "Arquitectura, flujo de eventos, comunicacion frontend-backend e integraciones Minecraft/TikTok"
    )
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    add_body(
        document,
        "Este informe se basa en la revision directa del codigo del repositorio, especialmente "
        "server/index.js, bridge/index.js, src/live-control.js, los hooks del dashboard, los "
        "componentes de overlay y la configuracion Electron. La aplicacion es una beta desktop/web "
        "para conectar TikTok LIVE con overlays, acciones automatizadas y juegos locales.",
    )

    add_heading(document, "Resumen ejecutivo", 1)
    add_body(
        document,
        "El proyecto tiene una idea clara: un panel React administra estado, acciones y triggers; "
        "un backend Node concentra API REST, WebSockets, TikTok, overlay, media y despachos; y un "
        "bridge local traduce eventos remotos hacia Minecraft, GTA/ChaosMod y otros procesos locales. "
        "La arquitectura funciona para una beta, pero el backend esta demasiado cargado y hay una "
        "inconsistencia critica entre el WebSocket que intenta usar el dashboard y el WebSocket que "
        "realmente expone el servidor.",
    )

    add_heading(document, "Arquitectura", 1)
    add_body(
        document,
        "La arquitectura real es hibrida: puede ejecutarse como app web de desarrollo, como backend "
        "Node que sirve dist/, o como app Electron que levanta backend y bridge como procesos hijo."
    )
    add_table(
        document,
        ["Componente", "Archivos principales", "Responsabilidad"],
        [
            [
                "Frontend React",
                "src/main.jsx, src/App.jsx, src/pages/DashboardApp.jsx, src/hooks/*",
                "Panel de control, CRUD de acciones/triggers, estado visual, links de overlay, pruebas manuales, subida de media y operaciones de TikTok/Spotify.",
            ],
            [
                "Overlay web",
                "src/components/overlay/OverlayScreens.jsx",
                "Pantallas para alertas, smart bar y song request. Consume REST inicial y WebSocket /ws/overlay; incluye polling de latest-event como respaldo.",
            ],
            [
                "Backend Node/Express",
                "server/index.js, server/state-store.js, server/media-*.js, server/spotify*.js",
                "API REST protegida, WebSocket hub, conexion TikTok, persistencia de estado, media, Spotify, overlay mirror, despacho de eventos y health checks.",
            ],
            [
                "Bridge local",
                "bridge/index.js, bridge/*.ps1",
                "Cliente WebSocket hacia backend publico/local; reexpone eventos en puertos locales, ejecuta RCON de Minecraft y dispara GTA/ChaosMod o GTAVWebhook.",
            ],
            [
                "Electron desktop",
                "electron/main.cjs, electron/preload.cjs",
                "Empaqueta la experiencia desktop, inicia backend y bridge, administra logs, importa cookies de TikTok y expone IPC seguro al renderer.",
            ],
            [
                "Modelo compartido",
                "src/live-control.js",
                "Defaults de estado, normalizadores, URLs, reglas de triggers, creacion de eventos de overlay y resumen de comandos.",
            ],
        ],
        widths=[1.35, 1.85, 3.3],
    )

    add_heading(document, "Flujo de eventos", 1)
    add_body(
        document,
        "El flujo central nace en TikTok o en una simulacion manual del panel. El backend normaliza el "
        "evento, actualiza historiales y metricas, evalua reglas y ejecuta acciones. El overlay recibe "
        "siempre un evento visual; Minecraft/GTA solo reciben payloads cuando la accion incluye esos outputs.",
    )
    add_table(
        document,
        ["Paso", "Detalle"],
        [
            ["1. Entrada", "TikTokLiveConnection emite CHAT, EMOTE, GIFT, FOLLOW, SHARE o LIKE; tambien existen /api/events/test y /api/actions/:id/test."],
            ["2. Normalizacion", "normalizeTikTokEvent convierte payloads externos en eventos internos con tipo, usuario, comentario, gift, emotes y timestamps."],
            ["3. Procesamiento", "processIncomingEvent registra emotes observados, guarda recentEvents, actualiza smart bar, evalua comandos de musica y chat mirror."],
            ["4. Matching", "matchesTrigger compara fuente, match text/gift/emote/audiencia y aplica cooldown por trigger."],
            ["5. Despacho", "dispatchAction genera overlayEvent, lo transmite a overlay/app, y si corresponde envia minecraft-command o gta-event."],
            ["6. Persistencia/estado", "StateStore guarda configuracion; recentEvents/recentDispatches viven en memoria; broadcastStatus actualiza estado operativo."],
        ],
        widths=[1.35, 5.15],
    )

    add_heading(document, "Comunicacion frontend-backend", 1)
    add_body(
        document,
        "La comunicacion usa REST para operaciones de control y WebSockets para actualizacion en tiempo real. "
        "Las rutas internas bajo /api quedan protegidas por la clave del panel, leida desde header "
        "X-Live-Control-Key o query key. El overlay usa una clave publica separada.",
    )
    add_table(
        document,
        ["Canal", "Uso", "Observaciones"],
        [
            ["REST /api/state", "Carga, guardado e importacion del estado del dashboard.", "El frontend tambien cachea en localStorage y compara updatedAt."],
            ["REST /api/tiktok/*", "Conectar/desconectar TikTok, sincronizar gifts/emotes, importar sesion desktop.", "Puede usar sessionid + tt-target-idc para sesion autenticada."],
            ["REST /api/media", "Subida/listado/borrado de archivos locales para overlay.", "Multer limita a 250 MB y los videos se normalizan para web."],
            ["REST /api/overlay/:slug", "Estado publico del overlay y ultimo evento.", "Incluye fallback por polling de latest-event."],
            ["WS /ws/overlay", "Estado y eventos para overlays.", "Protegido por overlayKey si esta configurada."],
            ["WS /ws/app", "Estado del panel y mensajes internos.", "El backend lo expone, pero el dashboard actualmente intenta /api/stream."],
            ["WS /ws/minecraft y /ws/gta", "Canales para bridge local.", "Protegidos por dashboardKey; tienen heartbeat y reconexion desde el bridge."],
        ],
        widths=[1.65, 2.25, 2.6],
    )

    add_heading(document, "Integracion Minecraft", 1)
    add_body(
        document,
        "Minecraft tiene dos caminos: ejecucion directa por RCON desde el backend y reenvio por bridge "
        "WebSocket para clientes/mods locales. Las acciones con output minecraft generan un payload "
        "minecraft-command a partir de buildBridgePayload. Si commandText existe, el backend intenta "
        "normalizarlo y ejecutarlo con rcon-client; el bridge tambien puede ejecutar RCON si useRcon esta activo."
    )
    add_bullets(
        document,
        [
            "Configuracion principal: profile.minecraftHost, minecraftPort y minecraftPassword en el estado; bridge-config.json para el bridge local.",
            "Puerto local por defecto del bridge Minecraft: ws://127.0.0.1:6135.",
            "Soporta modo generico y presets tipo Bedrock Box mediante minecraftMode, minecraftBedrockPresetId y minecraftBedrockPresetName.",
            "El chat mirror convierte comentarios de TikTok en comandos tellraw u otro formato configurado, con filtro para evitar espejar comandos del chat si esta activo.",
            "Punto a vigilar: hay responsabilidad duplicada de RCON entre backend y bridge; conviene elegir una autoridad unica para evitar doble ejecucion o resultados divergentes.",
        ],
    )

    add_heading(document, "Integracion TikTok", 1)
    add_body(
        document,
        "La integracion usa tiktok-live-connector. El panel llama /api/tiktok/connect con username y, opcionalmente, "
        "sessionid + tt-target-idc. En Electron, una ventana de login dedicada lee cookies de TikTok y las envia al "
        "backend por una ruta interna protegida con LIVE_CONTROL_DESKTOP_TOKEN.",
    )
    add_bullets(
        document,
        [
            "Eventos escuchados: chat, emote, gift, follow, share y like-burst.",
            "Los gifts tipo streak se filtran hasta repeatEnd para evitar ejecutar antes de que termine la repeticion.",
            "Al conectar, el backend intenta sincronizar el catalogo de gifts y emite mensajes de sistema al panel.",
            "Los emotes observados durante el live se incorporan al catalogo local y tambien existe sincronizacion autenticada.",
            "Riesgo base: tiktok-live-connector depende de reverse engineering y puede romperse si TikTok cambia protocolos, cookies o WebSocket.",
        ],
    )

    add_heading(document, "Puntos debiles", 1)
    add_table(
        document,
        ["Prioridad", "Punto debil", "Impacto"],
        [
            ["Alta", "Inconsistencia WebSocket del dashboard: frontend usa /api/stream y espera server-status/state-updated; backend expone /ws/app y emite status/state con payload.", "El panel puede perder actualizaciones en tiempo real y depender de recargas REST."],
            ["Alta", "server/index.js concentra demasiadas responsabilidades en un archivo de mas de 4.600 lineas.", "Dificulta pruebas, mantenimiento, aislamiento de errores y evolucion por dominios."],
            ["Alta", "Documentacion existente desfasada o contradictoria sobre ChaosMod/socket/HTTP/fallbacks.", "Aumenta el riesgo operativo y hace dificil diagnosticar fallos reales."],
            ["Media", "No se ven tests automatizados para contratos REST/WS, matching de triggers, RCON o payloads de bridge.", "Cambios pequenos pueden romper flujos de live sin aviso."],
            ["Media", "Secretos y claves se guardan en JSON/localStorage/AppData sin cifrado fuerte.", "Proteccion suficiente para beta local, debil para despliegues compartidos."],
            ["Media", "Despacho a bridge por WebSocket no tiene ack de ejecucion real.", "El backend sabe que envio a N clientes, pero no si Minecraft/GTA ejecuto correctamente."],
            ["Media", "RCON puede ejecutarse desde backend y bridge.", "Posible duplicacion de comandos o diferencias entre modo desktop y modo deploy."],
            ["Media", "Health checks de GTA mezclan configuracion, disponibilidad del bridge y prueba de endpoint local.", "Puede dar falsos negativos/positivos si el entorno no coincide con desktop."],
            ["Baja", "Logs con caracteres mojibake en varios archivos/documentos.", "Complica soporte y lectura profesional, aunque no rompe logica."],
            ["Baja", "Validacion de payloads y config hecha a mano.", "Errores de forma llegan tarde, durante ejecucion."],
        ],
        widths=[0.75, 2.65, 3.1],
    )

    add_heading(document, "Mejoras sugeridas", 1)
    add_table(
        document,
        ["Orden", "Mejora", "Resultado esperado"],
        [
            ["1", "Corregir el canal del dashboard: usar /ws/app en frontend o crear /api/stream compatible, y unificar nombres de mensajes.", "Tiempo real confiable para estado, dispatches, overlay y cambios de configuracion."],
            ["2", "Separar server/index.js en modulos: tiktok, dispatch, overlay, minecraft, music, auth, health, mirror, media.", "Menos riesgo por cambio y pruebas unitarias mas directas."],
            ["3", "Definir contratos de mensajes REST/WS con schemas compartidos y validacion runtime.", "Errores detectados antes de ejecutar acciones en vivo."],
            ["4", "Agregar ack/resultados del bridge para minecraft-command y gta-event.", "El panel puede mostrar enviado, recibido, ejecutado o fallido con causa."],
            ["5", "Elegir un solo responsable para RCON Minecraft, preferentemente el bridge local cuando el backend sea publico.", "Evita doble ejecucion y respeta la frontera local/remota."],
            ["6", "Crear tests de matchesTrigger, normalizeTikTokEvent, dispatchAction, buildBridgePayload y rutas criticas.", "Base segura para iterar sin romper el live."],
            ["7", "Mover secretos a almacenamiento seguro de Electron o variables de entorno, y no exponerlos en backups por defecto.", "Menor riesgo al compartir backups o logs."],
            ["8", "Actualizar documentacion tecnica eliminando secciones obsoletas y marcando version/fecha.", "Menos confusion entre comportamiento historico y actual."],
            ["9", "Agregar observabilidad: request ids, logs estructurados, niveles y panel de diagnostico.", "Soporte mas rapido cuando falla TikTok, bridge o ChaosMod."],
            ["10", "Endurecer media uploads: validacion MIME, extension, duracion, conversion asyncrona y cuotas.", "Menos riesgo de archivos incompatibles o pesados durante el live."],
        ],
        widths=[0.55, 3.05, 2.9],
    )

    add_heading(document, "Conclusion", 1)
    add_body(
        document,
        "El proyecto ya tiene una base funcional y bastante ambiciosa: TikTok LIVE entra como fuente de eventos, "
        "el backend decide que hacer, el overlay refleja la reaccion del stream y el bridge conecta con el mundo local "
        "de Minecraft/GTA. El proximo salto no deberia ser sumar mas features, sino estabilizar contratos, separar dominios "
        "y cerrar la brecha de WebSocket del dashboard. Con esas mejoras, la beta quedaria mucho mas predecible para uso real en vivo.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Live Control App - Analisis tecnico")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(85, 85, 85)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
